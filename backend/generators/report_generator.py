import re
import time
from xml.sax.saxutils import escape

from docx import Document as DocxDocument
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from backend.config import get_settings


def _slugify(title: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-")
    return f"{slug or 'report'}-{int(time.time())}"


def _references_lines(citations: list[dict]) -> list[str]:
    seen = set()
    lines = []
    for c in citations:
        label = c.get("url") or c.get("source") or "Unknown source"
        if label in seen:
            continue
        seen.add(label)
        page = f" (p. {c['page']})" if c.get("page") else ""
        lines.append(f"{c.get('source', label)}{page}" + (f" — {c['url']}" if c.get("url") else ""))
    return lines


def generate_report(title: str, sections: dict[str, str], citations: list[dict]) -> dict:
    settings = get_settings()
    reports_dir = settings.resolved_path(settings.reports_dir)
    slug = _slugify(title)
    references = _references_lines(citations)

    markdown_path = reports_dir / f"{slug}.md"
    docx_path = reports_dir / f"{slug}.docx"
    pdf_path = reports_dir / f"{slug}.pdf"

    _write_markdown(markdown_path, title, sections, references)
    _write_docx(docx_path, title, sections, references)
    _write_pdf(pdf_path, title, sections, references)

    return {
        "title": title,
        "markdown_path": str(markdown_path),
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path),
    }


def _write_markdown(path, title: str, sections: dict[str, str], references: list[str]) -> None:
    lines = [f"# {title}", ""]
    for section_title, content in sections.items():
        lines += [f"## {section_title}", "", content, ""]
    if references:
        lines += ["## References", ""] + [f"- {ref}" for ref in references]
    path.write_text("\n".join(lines), encoding="utf-8")


def _write_docx(path, title: str, sections: dict[str, str], references: list[str]) -> None:
    doc = DocxDocument()
    doc.add_heading(title, level=0)
    for section_title, content in sections.items():
        doc.add_heading(section_title, level=1)
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    if references:
        doc.add_heading("References", level=1)
        for ref in references:
            doc.add_paragraph(ref, style="List Bullet")
    for style_name in ("Normal",):
        doc.styles[style_name].font.size = Pt(11)
    doc.save(str(path))


def _write_pdf(path, title: str, sections: dict[str, str], references: list[str]) -> None:
    styles = getSampleStyleSheet()
    story = [Paragraph(escape(title), styles["Title"]), Spacer(1, 16)]
    for section_title, content in sections.items():
        story.append(Paragraph(escape(section_title), styles["Heading2"]))
        story.append(Spacer(1, 6))
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                story.append(Paragraph(escape(paragraph.strip()).replace("\n", "<br/>"), styles["BodyText"]))
                story.append(Spacer(1, 6))
    if references:
        story.append(Paragraph("References", styles["Heading2"]))
        for ref in references:
            story.append(Paragraph(f"• {escape(ref)}", styles["BodyText"]))
    SimpleDocTemplate(str(path), pagesize=LETTER).build(story)
